from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── 스타일 설정 ───
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '맑은 고딕'
    h.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

# ─── 헬퍼 함수 ───
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A478A')
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'EBF0F7')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run

# ═══════════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AW2026')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('2026 스마트공장·자동화산업전\n참가업체 심층 조사 보고서')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('2026년 3월 4~6일 | 서울 코엑스 전관\n24개국 500개 기업, 2,300부스, 약 80,000명')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 행사 개요
# ═══════════════════════════════════════════════════════════════
doc.add_heading('행사 개요', level=1)
add_table(doc,
    ['항목', '내용'],
    [
        ['행사명', 'AW 2026 (제36회 스마트공장·자동화산업전)'],
        ['일시', '2026년 3월 4일(수) ~ 6일(금), 3일간'],
        ['장소', '서울 코엑스(COEX) 전관 (A·B·C·D홀)'],
        ['규모', '24개국 500개 기업, 2,300부스, 약 80,000명 참관'],
        ['슬로건', '자율화, 지속가능성을 이끄는 힘 (Autonomy, The Driver of Sustainability)'],
        ['핵심 키워드', 'Physical AI, 휴머노이드, 자율제조(AX), AI 팩토리'],
        ['특별관', '스마트물류 특별관(D홀), AI 팩토리 특별관, Korea Vision Show(B홀)'],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. 로봇 그리퍼
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. 로봇 그리퍼 업체 및 대표 기술 (그리퍼 종류별)', level=1)

doc.add_heading('1-1. 그리퍼 종류별 비교', level=2)
add_table(doc,
    ['종류', '작동 원리', '가반하중', '장점', '단점', '적용 분야'],
    [
        ['공압', '압축 공기 피스톤', '0.1~수백kg', '단순, 저가, 고파지력', '2단계 제어, 공압라인', '자동차, 금속가공'],
        ['전동', '서보모터+볼스크류', '0.1~20kg', '정밀 제어, 저소음', '초기 비용 높음', '반도체, 전자부품'],
        ['진공', '흡착패드+진공발생기', '0.01~50+kg', '물체 손상 없음', '다공성 불가', '식품, 웨이퍼, 물류'],
        ['자석', '영구/전자석 자기력', '0.5~수백kg', '고온 대응', '강자성체만', '프레스, 철강'],
        ['소프트', '실리콘 공압 팽창', '0.01~3kg', '깨지기 쉬운 물체', '파지력 낮음', '식품, 제약'],
        ['다관절', '다관절 핑거 개별 모터', '0.5~5kg', '인간 수준 조작', '매우 고가', '연구, 휴머노이드'],
        ['적응형', '언더액추에이티드', '0.2~10kg', '다품종 대응', '최적 성능 낮음', '물류 피킹'],
        ['니들', '바늘 관통 파지', '0.05~5kg', '다공성 소재 대응', '구멍 발생', '섬유, 내장재'],
        ['툴체인저', '커플링 자동 교환', '1~2,500kg', '다공정 수행', '비용, 가반 차감', '용접, 조립'],
    ],
    col_widths=[2.2, 3.5, 2.2, 3, 2.8, 3]
)

doc.add_heading('1-2. AW2026 참가 확인 - 다관절/다지 그리퍼', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '핵심 스펙', '적용분야'],
    [
        ['테솔로 (TESOLLO)', '한국', 'Delto Gripper DG-5F', '20DoF 5지, 1.7kg, 파지력 20kg, 250Hz, 모듈형', '휴머노이드, 정밀조작'],
        ['테솔로', '한국', 'DG-5F-S (경량)', '20DoF, 1kg 미만, 고해상도 자기 엔코더', '경량 휴머노이드'],
        ['원익로보틱스', '한국', 'Allegro Hand V5 Sense', '16DoF 4지, 16개 압력센서, 촉각, ISO 9409', '연구, AI 학습'],
        ['원익로보틱스', '한국', 'Allegro Hand V5 4F Plus', '16DoF, 기존 3배 가반하중, CE 인증', '산업용 정밀파지'],
        ['티로보틱스', '한국', 'TR-WORKS 로봇 핸드', '3지/2지 교체형, 클린룸 대응', '반도체/디스플레이'],
        ['에이딘로보틱스', '한국', 'AIDIN Hand Gen2', '16DoF, 촉각센서(ATT), 시간당 1,000개', '물류, 휴머노이드'],
    ],
    col_widths=[3, 1.5, 4, 5, 3]
)

doc.add_heading('1-2. AW2026 참가 확인 - 전동/공압/AI센싱/적응형', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '그리퍼 종류', '핵심 스펙'],
    [
        ['씬그립 (SEENGRIP)', '한국', '옵티멈/이센셜/이코노미', '전동', '위치 인식, 힘센서 없이 부드러움 감지'],
        ['에이딘로보틱스', '한국', 'SUSGrip 스마트 그리퍼', '전동', 'BLDC 모터, 6축 F/T 센서 핑거팁'],
        ['주강로보테크', '한국', '산업용 공압 그리퍼', '공압', '37년 업력, 국내 대부분 로봇업체 공급'],
        ['모베이스전자', '한국', '택타일 센서 그리퍼', 'AI센싱', '프레셔 맵 실시간 시각화'],
        ['씨메스 (CMES)', '한국', '랜덤 팔레타이저', '적응형', '3D 비전+AI, 다양한 크기·재질 대응'],
    ],
    col_widths=[3.2, 1.5, 3.5, 2, 6]
)

doc.add_heading('1-3. 글로벌 주요 그리퍼 업체', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '그리퍼 종류', '특징'],
    [
        ['SCHUNK', '독일', 'EGA, PGN-plus-P', '공압/전동', '글로벌 리더, 3,000종+'],
        ['OnRobot', '덴마크', 'RG2, RG6, VGC10', '전동/진공/소프트', '협동로봇 플러그앤플레이'],
        ['Robotiq', '캐나다', '2F-85, Hand-E, 3-Finger', '전동/적응형', 'UR 호환, 적응형 선두'],
        ['Festo', '독일', 'FinGripper, DHEF', '공압/전동/소프트', '바이오닉 그리퍼 연구'],
        ['SMC', '일본', 'ZGS 진공, 전동/공압', '공압/전동/진공', '공압기기 세계 1위'],
        ['슈말츠 (Schmalz)', '독일', 'FQE, mGrip', '진공/소프트', 'Soft Robotics 인수, 식품'],
        ['Zimmer Group', '독일', '5000/6000/Magic', '공압/전동/진공', '40년+, IO-Link'],
        ['DESTACO', '미국', 'eRDH, RP/RA', '공압/전동', '파지력 40~2,000+lbs'],
        ['Piab', '스웨덴', 'piCOBOT Electric', '진공', '완전 전동식 진공'],
        ['Gimatic', '이탈리아', 'GS, MPXM', '공압/전동', '1,000만 사이클, FDA-H1'],
        ['SoftGripping', '독일', 'SoftGripper', '소프트', 'FDA 실리콘, 식품 전문'],
        ['Shadow Robot', '영국', 'Dexterous Hand', '다관절', '24DoF, 100+센서'],
    ],
    col_widths=[3, 1.8, 3.5, 3, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. 로봇 팔
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. 로봇 팔 업체 및 대표 기술 (로봇 암 / 협동로봇)', level=1)

doc.add_heading('2-1. 협동로봇 (Cobot) — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '핵심 스펙', '적용분야'],
    [
        ['화낙 (FANUC)', '일본', 'CRX-5iA~CRX-30iA', '6DoF, 5~30kg, 994~1889mm, ±0.03mm, NVIDIA AI', '제조, 용접, 팔레타이징'],
        ['유니버설 로봇 (UR)', '덴마크', 'UR3e~UR30', '6DoF, 3~30kg, 500~1750mm', '제조, 조립, 머신텐딩'],
        ['테크맨로봇', '대만', 'TM5~TM20', '6DoF, 4~20kg, 내장 비전+AI, DC타입, SEMI S2', '반도체, 물류'],
        ['뉴로메카', '한국', 'Indy 시리즈', '피지컬 AI, 30kg/50kg급 100% 국산', '제조, 물류, 서비스'],
        ['인아오리엔탈모터', '한국/일본', 'KOVR (3~5축)', '130×130mm 설치, Absolute Sensor', '반복 작업, 이송'],
    ],
    col_widths=[3, 2, 3, 5.5, 3]
)

doc.add_heading('2-2. AI 로봇 솔루션 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '핵심 스펙', '적용분야'],
    [
        ['씨메스 (CMES)', '한국', 'AI 디팔레타이징/피스피킹/용접', '3D 비전+AI, 티칭 불필요, 1,200회/h', '물류, 용접'],
        ['에이딘로보틱스', '한국', 'AIDIN Hand Gen2 + F/T 센서', '16DoF, One-stop Force-Aware', '정밀 파지'],
        ['원익로보틱스', '한국', 'WR-AMMR05-3F (AMMR)', '6DoF 팔+AMR, 도킹 ±1mm, 비전 AI', '반도체, 전자부품'],
    ],
    col_widths=[3, 1.5, 4, 5, 3]
)

doc.add_heading('2-3. 휴머노이드 로봇 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '핵심 스펙', '적용분야'],
    [
        ['보스턴 다이나믹스', '미국(현대차)', 'Atlas (전기형)', '190cm, 90kg, 56DoF, 50kg(순간), -20~40°C', '제조·물류 (미구동 국내 최초)'],
        ['티로보틱스', '한국', 'TR-WORKS', '180cm, 180kg, 20DoF, 클린룸', '반도체 (AW2026 첫 공개)'],
        ['현대차 로보틱스랩', '한국', 'MobED', '74×115cm, 4바퀴, CES 최고혁신상', '물류, 배송 (국내 최초)'],
        ['애지봇 (Agibot)', '중국', 'X2 / G2', 'X2: 130cm, 31DoF / G2: 49+DoF 바퀴형', '서비스 / 산업 (중국 빅5)'],
        ['유니트리 (Unitree)', '중국', 'G1', '127cm, 43DoF, 120Nm, 7km/h+, 3D 라이다', '제조·물류 (중국 빅5)'],
        ['푸리에 (Fourier)', '중국', 'GR-3C / GR-3', '165cm, 71kg, 55DoF', '산업 / 돌봄 (중국 빅5)'],
        ['레주 (Leju)', '중국', 'Kuavo 4 Pro', '~170cm, 26~30DoF, 360Nm, 화웨이 OS', '산업, 서비스 (중국 빅5)'],
    ],
    col_widths=[3, 2, 2.8, 5, 3.5]
)

doc.add_heading('2-4. 한국 대표 협동로봇 업체', level=2)
add_table(doc,
    ['업체명', '대표 제품', '특징'],
    [
        ['두산로보틱스', 'M/H/A 시리즈 (6~25kg)', '국내 점유율 1위, 글로벌 4위'],
        ['한화로보틱스', 'HCR 시리즈', '국내 빅3'],
        ['레인보우로보틱스', 'RB 시리즈, 이족보행', '삼성전자 파트너, 시가총액 6.9조원'],
        ['HD현대로보틱스', '산업용/협동로봇', '대형~협동 라인업 확장 중'],
    ],
    col_widths=[4, 5, 7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. AGV / AMR / AGF
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. AI 무인 자동화 이송 업체 및 대표 기술 (AGV / AMR / AGF)', level=1)

doc.add_heading('3-1. AGV vs AMR vs AGF 비교', level=2)
add_table(doc,
    ['구분', 'AGV', 'AMR', 'AGF'],
    [
        ['주행 방식', '고정경로 (자기테이프, QR)', '자율경로 (SLAM, LiDAR)', '자율/유도 하이브리드'],
        ['경로 변경', '인프라 물리 변경', 'SW 맵 즉시 변경', 'SW 맵 변경 가능'],
        ['적재용량', '수백kg ~ 수톤', '180kg ~ 1,500kg', '1톤 ~ 4톤+'],
        ['리프트 높이', '해당 없음', '해당 없음', '최대 11.5m'],
        ['도입 비용', '상대적 저렴', '중간 ~ 고가', '고가 (1억원~)'],
        ['대표 업체', '현대무벡스', '유진로봇, MiR, Geek+', 'BALYO, 러셀, VisionNav'],
    ],
    col_widths=[2.5, 4.5, 4.5, 4.5]
)

doc.add_heading('3-2. AMR 업체 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '적재용량', '내비게이션', '적용분야'],
    [
        ['유진로봇', '한국', 'GoCart 180~1500, 300 Omni', '180~1,500kg', '3D ToF LiDAR SLAM', '물류, 반도체'],
        ['현대무벡스', '한국', 'AMR 군집주행', '다양', 'SLAM, 군집관제', '물류, 제조'],
        ['현대글로비스', '한국', 'AMR+원키트 피킹', '-', '자율주행+AI 피킹', '자동차, 3PL'],
        ['현대차 로보틱스랩', '한국', 'MobED', '모듈형', 'AI+LiDAR+카메라', '물류, 안내'],
        ['긱플러스 (Geek+)', '중국', 'P800', '1,000kg', 'QR+자율주행', '이커머스'],
        ['빅웨이브로보틱스', '한국', '마로솔/솔링크', '-', '이기종 통합관제', '공항, 병원'],
    ],
    col_widths=[3, 1.5, 3.5, 2, 3, 3]
)

doc.add_heading('3-2. 유진로봇 GoCart 상세 사양', level=3)
add_table(doc,
    ['모델', '적재', '최대속도', '적재면', '안전인증'],
    [
        ['GoCart 180', '180kg', '1.0 m/s', '-', 'PL d'],
        ['GoCart 250', '250kg', '1.0 m/s', '-', 'PL d'],
        ['GoCart 500', '500kg', '1.0 m/s', '1222×822mm', 'PL d'],
        ['GoCart 1000', '1,000kg', '-', '-', 'PL d'],
        ['GoCart 1500', '1,500kg', '2.0 m/s', '1222×822mm', 'PL d'],
        ['GoCart 300 Omni', '300kg', '-', '-', 'PL d (옴니)'],
    ],
    col_widths=[3, 2.5, 2.5, 3.5, 3]
)

doc.add_heading('3-3. 한국 추가 AMR 업체', level=2)
add_table(doc,
    ['업체명', '대표 제품', '특징'],
    [
        ['트위니 (Twinny)', '나르고, 수다쟁이 나르고', '실내외 자율주행, 오더 피킹'],
        ['시스콘로보틱스', 'MPR(모바일 피킹 로봇)', '국내 최초 제조공정 AMR, 1,400대 공급'],
        ['티라로보틱스 (THIRA)', '2세대 AMR', 'LS 자회사, 경사/엘리베이터 대응'],
        ['마로로봇테크', '고중량 물류로봇', '올해의 로봇기업'],
        ['나우로보틱스', 'NUGO-P AMR (30kg)', '한양로보틱스 인수 통합'],
    ],
    col_widths=[4, 4.5, 7]
)

doc.add_heading('3-4. AGF (무인 지게차) — 한국 업체', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '특징'],
    [
        ['러셀로보틱스', '한국', 'AGF (기아 PV7)', '기아 40대 공급, 현대차·기아 AGF 82/84대'],
        ['에이로보틱스', '한국', 'AMR 포크리프트', 'LiDAR SLAM, ±10mm'],
        ['두산밥캣', '한국', 'AGF 카운터밸런스', '5개 센서, 원격 관제, 국내 지게차 1위'],
        ['모비우스', '한국', 'AFL', '기존 지게차 자율주행 업그레이드'],
        ['포테닛', '한국', 'AFL', '2017년부터 개발, 가상영역 안내'],
    ],
    col_widths=[3, 1.5, 3.5, 8]
)

doc.add_heading('3-4. AGF — 글로벌 업체', level=2)
add_table(doc,
    ['업체명', '국적', '대표 모델', '적재', '리프트', '특징'],
    [
        ['비전나비 (VisionNav)', '중국', 'VNR16, VNE30~40', '1.4~4톤', '11.5m', '비전 SLAM, 한국법인'],
        ['멀티웨이 (Multiway)', '중국', 'MW-SE15, R16', '1.5~2톤', '3m+', '레이저 SLAM, 500대 기업'],
        ['BALYO', '프랑스', 'VEENY, REACHY', '다양', '17m', 'ABI 1위, KION/Amazon'],
        ['Seegrid', '미국', '팔레트 트럭/터그', '중량물', '-', '컴퓨터 비전 내비게이션'],
        ['Jungheinrich', '독일', '자율주행 지게차', '다양', '다양', '유럽 3대, Rocrich 합작'],
        ['Toyota Industries', '일본', 'Material Handling', '다양', '다양', '세계 1위 지게차'],
    ],
    col_widths=[3, 1.3, 2.8, 1.8, 1.5, 5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. AI 물류 자동화
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. AI 물류 자동화 업체 및 대표 기술', level=1)

doc.add_heading('4-1. 토털 물류 자동화 솔루션 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품/플랫폼', '핵심 기술', '주요 고객'],
    [
        ['현대글로비스', '한국', '오르카(ORCA) WCS, 팔레트 셔틀, 아틀라스', '피지컬 AI, 크로스플랫폼 AMR 통합', '에코프로, 현대차그룹'],
        ['현대무벡스', '한국', 'SRM, 셔틀, AS/RS, 옴니소터, 3D DT', 'AI 군집, 웹 원격 관제', 'SK 테네시, LG화학, 한국타이어'],
        ['CJ올리브네트웍스', '한국', 'AI 비전, 자율운전, 에이전틱 AI', 'MES/RTDB 기반 에이전틱 AI', 'CJ그룹, 다쏘/지멘스'],
        ['엠투아이+에스엠코어', '한국', '스마트팩토리+물류자동화', 'AI 데이터 의사결정', 'SK하이닉스 (수주 2,000억+)'],
    ],
    col_widths=[3, 1.3, 4, 4, 4]
)

doc.add_heading('4-2. 자동 창고 / 분류 / AI 피킹 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '분류', '대표 제품', '핵심 기술', '처리 능력'],
    [
        ['긱플러스', '자동창고', 'SkyCube, PopPick', 'AI 저장밀도 4배 최적화', '650토트/h'],
        ['아세테크', '자동창고', '오토스토어, neXos', '엔드투엔드 SI', '풀 자동화 라인'],
        ['TXR로보틱스', '분류', '휠소터, 싱귤레이터, 셀루베이어', '300개 모터 개별 제어', '싱귤레이터 6,000+/h'],
        ['미르 (MiR)', '분류', '린콘(LINCON), iTS', '리니어 모터+자기부상', '기존 5배 속도'],
        ['씨메스', 'AI 피킹', '랜덤 팔레타이징, 피스피킹', '3D 비전+AI, 무작위 인식', '피스피킹 1,200회/h'],
    ],
    col_widths=[2.5, 2, 4, 4, 3.5]
)

doc.add_heading('4-3. WCS/관제/디지털트윈 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '분류', '대표 제품', '핵심 기술'],
    [
        ['현대글로비스', 'WCS', '오르카(ORCA)', '창고 분석→최적 배치, 다기종 AMR 통합'],
        ['빅웨이브로보틱스', '관제', '솔링크(SOLlink), 마로솔', '이기종 로봇 통합(GS 1등급), 5,000건 DB'],
        ['현대무벡스', '디지털트윈', '3D 디지털 트윈', 'AS/RS+AMR 실시간, 웹 원격'],
        ['와고코리아', '디지털트윈', 'DT 통합 제어·시각화', 'ctrlX CORE 연동'],
    ],
    col_widths=[3, 2.5, 4, 7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. AI 선별 자동화
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. AI 선별 자동화 업체 및 대표 기술', level=1)

doc.add_heading('5-1. 딥러닝/AI 기반 품질검사·선별 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', 'AI 기술', '정확도/성능', '적용 분야'],
    [
        ['세이지 (SAIGE)', '한국', 'SAIGE Vision, AI 에이전트', '비지도학습, MLOps', '소량 데이터 적용', '반도체, PCB'],
        ['뉴로클', '한국', 'Neuro-T/R/X, Engine', '오토딥러닝, NPU', '비전문가 고정밀', '제조, 의료(25개국)'],
        ['라온피플', '한국', 'ADC, Hi FENN, ODIN', 'LLM+RAG, VLM', '불량률 20~30%↓', '반도체, 건설, 물류'],
        ['슈퍼브에이아이', '한국', 'ZERO 파운데이션 모델', '제로샷(Zero-shot)', '라벨링 불필요', '제조, 안전, 물류'],
        ['브이원테크', '한국', 'withAI, Edge Device', '자체 딥러닝', '미세 결함, 고속', '인라인 검사'],
        ['마키나락스', '한국', 'Weld VisionX, Runway', '3D AI 비전', '제로 티칭 용접', '조선, 중공업'],
        ['디플리 (Deeply)', '한국', 'Listen AI (음향 AI)', '음향 딥러닝', '99.78%+, 1초 미만', '모터, 자동차부품'],
        ['CJ올리브네트웍스', '한국', '비전 AI, 에이전틱 AI', '딥러닝+에이전틱', '실시간 이상 판단', '식품, 화학'],
        ['피아이이 (PIE)', '한국', 'AI 비전+NDT', '딥러닝+비파괴', '비전+NDT 통합', '자동차, CNC'],
    ],
    col_widths=[2.5, 1.2, 3, 2.5, 2.5, 2.5]
)

doc.add_heading('5-2. 초분광 AI / 특수 선별 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', 'AI 기술', '핵심 스펙', '적용 분야'],
    [
        ['오즈레이 (OZ Ray)', '한국', 'SWIR 초분광 카메라', '초분광+AI', '0.9~1.7μm, 256밴드', '재질, 이물질, 식품'],
        ['비케이인스트루먼트', '한국', 'RESONON/ClydeHSI', '초분광(UV~NIR)', '수백 채널 분할', '식품, 이차전지'],
        ['엘로이랩', '한국', '초분광 AI 자동화', '초분광+딥러닝+에어젯', '수백 대역 분석', '식품 이물질, 소재'],
        ['하이퍼그램', '한국', 'HG VNIR Pro', '압축식 초분광(세계최초)', '스캔프리 풀프레임', '산업검사, 소재'],
    ],
    col_widths=[3, 1.3, 3, 3, 3, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. 비전 선별
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. 비전 선별 업체 및 대표 기술 (AI 비전 / 머신비전 / 스마트 비전)', level=1)

doc.add_heading('6-1. 머신비전 카메라/센서 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '핵심 스펙', '적용 분야'],
    [
        ['Basler', '독일', 'MDFI, TDI, DT 비전', 'TDI 라인스캔, 디지털트윈 시뮬레이션', '철도, 반도체'],
        ['비투에스 (B2S)', '한국', 'VIENEX IV CIS, 3D, AFTS', 'CIS 일체형, AFTS 1μm 초점, CoF', '반도체, 이차전지'],
        ['싸이로드 (Cylod)', '한국', 'SICK 3D, Triton SWIR', '3D 마이크론/kHz, SWIR IP67', '정밀측정, 외관'],
        ['LUCID Vision Labs', '-', '3D ToF, 25GigE, SWIR', '고정밀 3D, 초고속', '로봇 가이딩'],
        ['뷰웍스', '한국', '초고해상도/라인스캔', '에어리어/라인 풀라인업', '반도체, 디스플레이'],
        ['Goermicro', '중국', 'CW-iToF, RGBD', '3D ToF 산업용 (부스 B706)', '스마트 제조'],
        ['넥센서', '한국', '3D 광학 계측', 'Advanced Package 특화 (B744)', '반도체'],
    ],
    col_widths=[3, 1.3, 3.5, 5, 3]
)

doc.add_heading('6-2. 비전 조명/광학 & 엣지 AI — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '핵심 스펙', '분류'],
    [
        ['에드몬드옵틱스', '미국', 'SilverTL 텔레센트릭', 'SWIR 900~1700nm, Bi-Telecentric', '조명/광학'],
        ['비투에스', '한국', 'UV~SWIR 필터, Rod Lens', '전 대역, 난반사 제거', '조명/광학'],
        ['젝스컴퍼니', '한국', 'JECS-1400GB', '인텔 울트라 NPU, GPU 불필요', '엣지 AI'],
        ['에이수스', '대만', 'PE3000N, NUC 16 Pro', 'Jetson, 7.5배 성능', '엣지 AI'],
        ['브이원테크', '한국', 'Edge Device', '로봇암 장착, 인라인 추론', '엣지 AI'],
        ['뉴로클', '한국', 'Neuro-T Engine', 'GUI 없이 임베딩, NPU', '엣지 AI'],
    ],
    col_widths=[3, 1.3, 3.5, 5, 2.5]
)

doc.add_heading('6-3. 비전+로봇+AI 통합 솔루션 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '핵심 기술', '적용 분야'],
    [
        ['테크맨로봇', '대만', 'TM AI Cobot', '내장 비전+AI 원플랫폼', '반도체, 물류'],
        ['씨메스', '한국', '팔레타이저/피스피킹/가이던스', '3D 비전+AI 피지컬 AI', '물류, 제조'],
        ['마키나락스', '한국', 'Weld VisionX', '3D AI 비전 제로 티칭 용접', '조선, 중공업'],
        ['피아이이', '한국', 'CNC 핸들러/헤드램프 검사', '비전+로봇+NDT 통합', '자동차, CNC'],
        ['비투에스', '한국', '직관(JIKGWAN)', '카메라+조명+필터+3D 원스톱', '비전 시스템 구축'],
    ],
    col_widths=[3, 1.3, 4, 4.5, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. AI PLC 통합 솔루션
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. AI PLC 통합 솔루션 업체 및 대표 기술', level=1)

doc.add_heading('7-1. AI PLC 제조사 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', 'AI 기능', '통신 프로토콜', '특징'],
    [
        ['LS일렉트릭', '한국', 'SU-CM70', 'SW정의, AI팩토리 연동', 'EtherNet/IP, EtherCAT', '최대 270m², 30부스'],
        ['로크웰', '미국', 'ControlLogix 5590', 'AI자율생산, Copilot', 'EtherNet/IP, CIP', '5개 존 구성'],
        ['싸이몬', '한국', 'CICON', 'AI Copilot, Code Converter', 'OPC UA, Modbus', '인지형 자동화'],
        ['파나소닉', '일본', 'FP7 PLC', 'FA 디바이스', 'EtherCAT, RTEX', '고속·고정밀'],
    ],
    col_widths=[2.5, 1.2, 3, 3, 3, 3.5]
)

doc.add_heading('7-2. AI 기반 SCADA/HMI 플랫폼 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', 'AI 기능'],
    [
        ['LS일렉트릭', '한국', 'AI팩토리 패키지, LS SHE, 블랙박스, LLM 진단', '비전AI 안전관제, LLM 자연어 진단, 예지보전'],
        ['슈나이더 일렉트릭', '프랑스', 'EcoStruxure, AVEVA, ETAP, PME', 'AI 예측·최적화, 분산 AI 통합분석'],
        ['싸이몬', '한국', 'SCADA PRO', 'AI Assistant: 자연어→분석, 스크립트/3D 자동 생성'],
        ['오토닉스', '한국', 'SCADAMaster, iTP HMI', 'SCADA 스마트 품질관리, AMR 제어'],
        ['엠투아이+에스엠코어', '한국', '스마트팩토리 플랫폼', 'AI 자율제조: 제어+공정+물류 통합'],
    ],
    col_widths=[3.5, 1.5, 5, 6]
)

doc.add_heading('7-3. 개방형 자동화 플랫폼 (SW 정의, 멀티벤더)', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', '핵심 기술'],
    [
        ['슈나이더 일렉트릭', '프랑스', 'EAE (개방형 SW 정의)', '멀티벤더 호환, 단계적 DX'],
        ['보쉬렉스로스', '독일', 'ctrlX AUTOMATION, ctrlX CORE', '앱 기반 확장, 멀티벤더'],
        ['LS일렉트릭', '한국', 'SU-CM70 (SW정의)', '기능 추가·확장 가능'],
        ['로크웰', '미국', 'FactoryTalk Design Studio, Logix Echo', 'Copilot, 에뮬레이션'],
    ],
    col_widths=[3.5, 1.5, 5, 6]
)

doc.add_heading('7-4. 예지보전/설비진단 & OT 보안 — AW2026 참가 확인', level=2)
add_table(doc,
    ['업체명', '국적', '대표 제품', 'AI 기능', '분류'],
    [
        ['LS일렉트릭', '한국', '블랙박스, LLM 진단', '이상 기록·분석, 자연어 진단', '예지보전'],
        ['마키나락스', '한국', 'Runway, DrawX, Weld VisionX', 'GPU 동적분할, 제로티칭', '예지보전'],
        ['디플리', '한국', 'Listen AI', '음향 AI 99.78%+', '예지보전'],
        ['두산 DI BU', '한국', 'AI 예지정비', 'Predictive Maintenance', '예지보전'],
        ['포스코DX', '한국', 'AX 융합 기술', 'AI 제조 혁신', '예지보전'],
        ['로크웰', '미국', 'SecureOT', 'OT 사이버보안 플랫폼', 'OT 보안'],
        ['힐셔 (Hilscher)', '독일', 'netX 900', '멀티프로토콜+보안 내장', 'OT 보안'],
    ],
    col_widths=[3, 1.3, 3.5, 4, 2.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 주요 트렌드
# ═══════════════════════════════════════════════════════════════
doc.add_heading('주요 트렌드 요약', level=1)
add_table(doc,
    ['트렌드', '내용'],
    [
        ['Physical AI / 자율제조(AX)', '올해 핵심 테마. 로봇이 물리 환경 인식→자율 판단·동작'],
        ['휴머노이드 대거 등장', '아틀라스, TR-WORKS, 에이르, 중국 빅5(애지봇/유니트리/푸리에/레주/화웨이)'],
        ['SW 정의 자동화(SDA)', 'LS SU-CM70, 슈나이더 EAE, 보쉬 ctrlX — HW 아닌 SW로 기능 정의'],
        ['LLM/자연어 AI 통합', 'LS(LLM 진단), 싸이몬(Assistant), 로크웰(Copilot), 벡호프(CoAgent)'],
        ['파운데이션 모델/제로샷', '슈퍼브AI ZERO, 세이지 — 라벨링 없이 즉시 비전 AI 적용'],
        ['음향 AI 품질검사', '디플리 Listen AI — 99.78%+ 정확도, 비(非)비전 선별 부상'],
        ['초분광 AI', '오즈레이, 비케이, 엘로이랩 — 화학적 차이 감지 선별'],
        ['AGV→AMR 전환 가속', 'SLAM 기반 자율경로, 군집주행, 이기종 통합 관제'],
        ['AGF 시장 급성장', 'LiDAR+SLAM 무인지게차, 기존 지게차 업그레이드 방식'],
        ['개방형 멀티벤더 호환', '슈나이더, 보쉬, LS — 타사 설비 호환 개방형 아키텍처'],
    ],
    col_widths=[4, 12]
)

# ─── 저장 ───
output_path = os.path.join(r'C:\Users\모재민\Documents\sample', 'AW2026_참가업체_조사보고서.docx')
doc.save(output_path)
print(f'저장 완료: {output_path}')
